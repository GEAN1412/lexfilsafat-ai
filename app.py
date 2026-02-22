import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO
import pandas as pd
import os
from datetime import datetime
import json
from PIL import Image, ImageDraw, ImageFont
import textwrap
import yfinance as yf

# ==========================================
# 1. KONFIGURASI HALAMAN
# ==========================================
st.set_page_config(page_title="LexFilsafat AI - Super App", page_icon="⚖️", layout="wide")

# KONFIGURASI DESAIN INSTAGRAM
FONT_HEADLINE_PATH = "Roboto-Bold.ttf"
FONT_BODY_PATH = "Roboto-Regular.ttf"
BG_COLOR = "#1E1E1E"
TEXT_COLOR = "#FFFFFF"
ACCENT_COLOR = "#FFD700"

# ==========================================
# 2. SETUP API KEY (AMAN DARI LEAK)
# ==========================================
try:
    API_KEY = st.secrets["GEMINI_API_KEY"]
except:
    API_KEY = ""

if API_KEY == "":
    st.error("🚨 Sistem Terkunci: API Key Gemini belum dikonfigurasi di Secrets.")
    st.stop()

genai.configure(api_key=API_KEY)
model = genai.GenerativeModel('gemini-2.5-flash')

# ==========================================
# FUNGSI BANTUAN (Word, Database, Gambar IG)
# ==========================================
def create_word_docx(teks_analisis, judul_perkara):
    doc = Document()
    doc.add_heading(f'Dokumen Hukum: {judul_perkara}', 0)
    doc.add_paragraph(teks_analisis)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def simpan_ke_database(nama, email, kasus):
    file_path = 'database_perkara.csv'
    data_baru = pd.DataFrame({'Tanggal': [datetime.now().strftime("%Y-%m-%d %H:%M:%S")], 'Nama': [nama], 'Email': [email], 'Kasus': [kasus]})
    if os.path.exists(file_path):
        data_lama = pd.read_csv(file_path)
        data_final = pd.concat([data_lama, data_baru], ignore_index=True)
    else:
        data_final = data_baru
    data_final.to_csv(file_path, index=False)

def create_instagram_slide(headline, body, slide_num):
    img = Image.new('RGB', (1080, 1080), color=BG_COLOR)
    draw = ImageDraw.Draw(img)
    try:
        font_h = ImageFont.truetype(FONT_HEADLINE_PATH, 80)
        font_b = ImageFont.truetype(FONT_BODY_PATH, 45)
        font_s = ImageFont.truetype(FONT_BODY_PATH, 30)
    except:
        font_h = ImageFont.load_default()
        font_b = ImageFont.load_default()
        font_s = ImageFont.load_default()

    margin = 80
    wrapped_h = textwrap.fill(headline, width=20)
    draw.text((margin, 150), wrapped_h, font=font_h, fill=ACCENT_COLOR)
    bbox = draw.textbbox((margin, 150), wrapped_h, font=font_h)
    
    wrapped_b = textwrap.fill(body, width=35)
    draw.text((margin, bbox[3] + 60), wrapped_b, font=font_b, fill=TEXT_COLOR)
    
    draw.rectangle([(margin, 980), (1080-margin, 990)], fill=ACCENT_COLOR)
    draw.text((margin, 1000), f"Slide {slide_num} • LexFilsafat AI", font=font_s, fill=TEXT_COLOR)
    return img

# ==========================================
# 3. SIDEBAR NAVIGASI
# ==========================================
st.sidebar.title("⚖️ LexFilsafat Menu")
menu = st.sidebar.radio("Pilih Layanan:", 
    ["Analisis Hukum Umum", "Radar Investasi & Hukum", "Dashboard Admin 🔒"]
)

# ==========================================
# MENU 1: ANALISIS UMUM
# ==========================================
if menu == "Analisis Umum":
    st.title("⚖️ Analisis Perkara Hukum Umum")
    user_input = st.text_area("Deskripsikan kronologi perkara atau istilah hukum:", height=150)
    
    st.info("💡 **Fitur Premium:** Masukkan nama dan email/WhatsApp Anda untuk mengunduh Draft Dokumen Hukum yang mendetail.")
    col1, col2 = st.columns(2)
    with col1: user_nama = st.text_input("Nama Anda (Opsional):")
    with col2: user_email = st.text_input("Email / WhatsApp (Opsional):")

    if st.button("Analisis Kasus"):
        if not user_input.strip():
            st.warning("⚠️ Isi dulu kronologi perkaranya.")
        else:
            with st.spinner("Menganalisis..."):
                try:
                    prompt = f"Sebagai Ahli Hukum. Analisis: {user_input}. Berikan Kualifikasi, Pasal, Filsafat, Rujukan."
                    hasil = model.generate_content(prompt).text
                    st.success("Selesai!")
                    st.markdown(hasil)
                    
                    if user_nama and user_email:
                        simpan_ke_database(user_nama, user_email, user_input)
                        prompt_draft = f"Buatkan draft surat hukum formal untuk: {user_input}"
                        draft = model.generate_content(prompt_draft).text
                        docx_file = create_word_docx(draft, "Draft_Hukum")
                        st.download_button("📥 Unduh Draft (.docx)", data=docx_file, file_name="Draft_Premium.docx")
                except Exception as e:
                    st.error(f"Error AI: {e}")

# ==========================================
# MENU 2: RADAR INVESTASI & HUKUM (yfinance)
# ==========================================
elif menu == "Radar Investasi & Hukum":
    st.title("📈 Radar Investasi & Risiko Hukum")
    
    tab_saham, tab_lainnya = st.tabs(["Saham (BEI)", "Reksadana, Obligasi & Crypto"])
    
    with tab_saham:
        st.subheader("Analisis Emiten Bursa Efek Indonesia")
        ticker_input = st.text_input("Masukkan Kode Emiten (4 Huruf):", placeholder="Contoh: BBCA, SIDO, atau GOTO")
        
        if st.button("Analisis Emiten 🚀"):
            if not ticker_input:
                st.warning("⚠️ Masukkan kode emiten terlebih dahulu!")
            else:
                with st.spinner("Menarik data pasar & meracik opini legal..."):
                    try:
                        ticker_symbol = f"{ticker_input.upper()}.JK"
                        stock = yf.Ticker(ticker_symbol)
                        hist = stock.history(period="5y")
                        info = stock.info
                        
                        if hist.empty:
                            st.error("Data emiten tidak ditemukan. Pastikan kode benar.")
                        else:
                            st.markdown(f"### Grafik Harga 5 Tahun ({ticker_input.upper()})")
                            st.line_chart(hist['Close'], color="#FFD700")
                            
                            st.markdown("### Indikator Fundamental")
                            col1, col2, col3, col4 = st.columns(4)
                            col1.metric("Harga Saat Ini", f"Rp {hist['Close'].iloc[-1]:,.0f}")
                            col2.metric("P/E Ratio (PER)", round(info.get('trailingPE', 0), 2) if info.get('trailingPE') else "N/A")
                            col3.metric("P/B Ratio (PBV)", round(info.get('priceToBook', 0), 2) if info.get('priceToBook') else "N/A")
                            col4.metric("Market Cap", f"Rp {info.get('marketCap', 0) / 1e12:,.2f} T" if info.get('marketCap') else "N/A")
                            
                            prompt_emiten = f"""
                            Sebagai Corporate Lawyer, berikan ulasan emiten {ticker_input.upper()} di BEI.
                            1. Profil Bisnis Singkat.
                            2. Risiko Hukum Sektoral.
                            3. Perlindungan Investor Ritel menurut UU Pasar Modal.
                            """
                            st.markdown(model.generate_content(prompt_emiten).text)
                    except Exception as e:
                        st.error(f"Error data pasar: {e}")
                        
    with tab_lainnya:
        st.subheader("Analisis Instrumen Lainnya")
        aset_input = st.text_area("Sebutkan instrumen yang ingin dianalisis:")
        if st.button("Analisis Aset"):
            if aset_input:
                with st.spinner("Menganalisis regulasi..."):
                    prompt_aset = f"Analisis hukum investasi ini di Indonesia: {aset_input}. Jelaskan regulasi Bappebti/OJK."
                    st.markdown(model.generate_content(prompt_aset).text)

# ==========================================
# MENU 3: DASHBOARD ADMIN (PASSWORD PROTECTED)
# ==========================================
elif menu == "Dashboard Admin 🔒":
    st.title("🔒 Area Khusus Admin")
    password = st.text_input("Password:", type="password")
    
    if password == "lexai1234":
        st.success("Akses Admin Diterima.")
        st.markdown("---")
        
        tab1, tab2 = st.tabs(["🎥 Generator Konten (Auto-Image)", "📊 Database Leads"])
        
        with tab1:
            st.subheader("Dapur Konten Otomatis")
            ide_konten = st.text_input("Topik Konten:", placeholder="Contoh: Pinjol Ilegal")
            platform = st.selectbox("Format:", [
                "Instagram Feed (Auto-Generate Gambar)",
                "YouTube Shorts / Veo (Naskah Video)"
            ])
            
            if st.button("Generate Konten 🚀"):
                if not ide_konten.strip():
                    st.warning("Masukkan topik dulu!")
                else:
                    if "Instagram" in platform:
                        with st.spinner("Meracik naskah & Menggambar slide..."):
                            try:
                                prompt_ig = f"""
                                Kamu adalah Social Media Manager. Buat konten Instagram 5 slide tentang: "{ide_konten}".
                                PENTING: Output HANYA format JSON murni.
                                {{
                                  "caption": "Caption lengkap...",
                                  "slides": [
                                    {{"headline": "Judul 1", "body": "Isi..."}},
                                    {{"headline": "Judul 2", "body": "Isi..."}},
                                    {{"headline": "Judul 3", "body": "Isi..."}},
                                    {{"headline": "Judul 4", "body": "Isi..."}},
                                    {{"headline": "Judul 5", "body": "Isi..."}}
                                  ]
                                }}
                                """
                                response = model.generate_content(prompt_ig)
                                json_str = response.text.replace("```json", "").replace("```", "").strip()
                                data = json.loads(json_str)
                                
                                st.subheader("1. Caption")
                                st.code(data['caption'], language='text')
                                
                                st.subheader("2. Slide Gambar")
                                cols = st.columns(5)
                                for i, slide in enumerate(data['slides']):
                                    img = create_instagram_slide(slide['headline'], slide['body'], i+1)
                                    with cols[i]:
                                        st.image(img, caption=f"Slide {i+1}", use_column_width=True)
                            except Exception as e:
                                st.error(f"Gagal generate gambar: {e}")
                    else:
                        with st.spinner("Menulis naskah video..."):
                            prompt_video = f"""
                            Buatkan naskah video Shorts tentang: "{ide_konten}".
                            Format Markdown Table:
                            | No | Naskah Voiceover | Visual Prompt (English for Veo) |
                            |---|---|---|
                            | 1 | ... | `cinematic shot...` |
                            """
                            st.markdown(model.generate_content(prompt_video).text)

        with tab2:
            st.subheader("Data Leads")
            if os.path.exists('database_perkara.csv'):
                df = pd.read_csv('database_perkara.csv')
                st.dataframe(df)
            else:
                st.info("Belum ada data.")

    elif password != "":
        st.error("Password Salah!")

# ==========================================
# FOOTER / CREDIT EKSKLUSIF
# ==========================================
st.sidebar.markdown("---")
st.sidebar.caption("Dibuat oleh Gean Pratama Adiaksa SH")
st.markdown("---")
st.caption("© 2026 | Dibuat oleh Gean Pratama Adiaksa SH - Analisis Hukum & Ekonomi")

